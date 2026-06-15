from docx import Document
import os


def create_docx(text: str, filename: str) -> str:
    """
    Creates a .docx file with the provided text.
    Returns the path to the file.
    """
    doc = Document()
    # Split the text into paragraphs for better formatting
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph()

    # Ensure the directory exists
    os.makedirs("data", exist_ok=True)
    file_path = os.path.join("data", filename)
    doc.save(file_path)
    return file_path
